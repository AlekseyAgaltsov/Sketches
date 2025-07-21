import pandas as pd
from docx import Document
import re
# pip install openpyxl

# Чтение Excel-файла
file_path = 'Журнал.xlsx'
df = pd.read_excel(file_path)

numb = int(input('Ввод номера заключения: '))

# Номер строки (начиная с 1)
row_number = numb - (11945 + 7)

# Список русских названий месяцев
months = [
    'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
    'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
]

# Получаем дату из DataFrame
dt = df.iloc[row_number - 1].iloc[1]
day = dt.day
month_name = months[dt.month - 1]  # получаем правильное склонённое название месяца
month = dt.strftime('%m')
year = dt.strftime('%y')
yearf = dt.strftime('%Y')

# Построение строки с датой
z_data = f'{day} {month_name} {yearf}'
zf_data = f'{day}.{month}.{year}'

# Остальные данные из DataFrame
zakl_nomer = int(df.iloc[row_number - 1].iloc[0])
zname = df.iloc[row_number - 1].iloc[2]
zok = df.iloc[row_number - 1].iloc[3]
zokname = df.iloc[row_number - 1].iloc[4]
zwho = df.iloc[row_number - 1].iloc[6]
ztp = df.iloc[row_number - 1].iloc[8]
zt = df.iloc[row_number - 1].iloc[9]
zm = df.iloc[row_number - 1].iloc[10]
zn = df.iloc[row_number - 1].iloc[12]

# Переменные для замены
variables = {
    '{zakl_nomer}': str(zakl_nomer),
    '{z_data}': z_data,  # готовая строка даты
    '{zf_data}': zf_data,  # готовая строка даты
    '{zname}': str(zname),
    '{zok}': str(zok),
    '{zokname}': str(zokname),
    '{zwho}': str(zwho),
    '{ztp}': str(ztp),
    '{zt}': str(zt),
    '{zm}': str(zm),
    '{zn}': str(zn)
}

# Открываем шаблон документа Word
doc = Document('Заключение ВИК.docx')

# Обрабатываем каждую переменную в каждом параграфе документа
pattern = '|'.join(map(re.escape, variables.keys()))


# Функция для полной замены переменных в тексте
def replace_all_variables(text):
    return re.sub(pattern, lambda match: variables.get(match.group(), match.group()), text)


# Полностью перебираем весь документ и делаем замену
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Собираем полные строки и меняем их разом
                full_text = ''.join(run.text for run in paragraph.runs)

                # Замещаем все переменные в полном тексте абзаца
                replaced_text = replace_all_variables(full_text)

                # Очищаем старый текст и вставляем заново
                for run in paragraph.runs[:]:
                    paragraph._p.remove(run._r)

                # Записываем обратно готовый текст
                paragraph.add_run(replaced_text)

# Сохраняем изменения
doc.save(f'Заключение_{zakl_nomer}.docx')